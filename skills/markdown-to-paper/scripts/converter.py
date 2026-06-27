#!/usr/bin/env python3
"""
Main converter: Markdown → Word (.docx) or PDF (IEEE Transactions format)
Supports custom reference templates for Word and .tex/.cls for PDF.
Default PDF template: IEEE Transactions (IEEEtran.cls, two-column)
"""

import argparse
import os
import re
import sys
import subprocess
import tempfile
import shutil
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from markdown_parser import parse_markdown


# ─────────────────────────────────────────────────────────────────────────────
# Word Document Generation (python-docx)
# ─────────────────────────────────────────────────────────────────────────────

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def set_cell_border_threeline(cell):
    """Apply three-line table borders to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "left", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_threeline_table(doc, table_data):
    """Add a three-line academic table to the Word document."""
    headers = table_data["headers"]
    rows = table_data["rows"]
    caption = table_data.get("caption", "")

    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after = Pt(2)
        run = cp.add_run(caption)
        run.bold = True
        run.font.size = Pt(9)

    n_cols = len(headers)
    table = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    table.style = "Table Grid"

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border_threeline(cell)

    for r_idx, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < len(table.rows[r_idx + 1].cells):
                cell = table.rows[r_idx + 1].cells[j]
                cell.text = cell_text
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                set_cell_border_threeline(cell)

    for j in range(n_cols):
        tc = table.rows[0].cells[j]._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "bottom"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "double")
            border.set(qn("w:sz"), "6")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    doc.add_paragraph()


def render_text_paragraph(doc, text):
    """Add a text paragraph with basic markdown formatting."""
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    segments = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for seg in segments:
        if seg.startswith("**") and seg.endswith("**"):
            run = p.add_run(seg[2:-2])
            run.bold = True
        elif seg.startswith("*") and seg.endswith("*") and not seg.startswith("**"):
            run = p.add_run(seg[1:-1])
            run.italic = True
        else:
            p.add_run(seg)


def markdown_to_docx(md_text, output_path, ref_template=None):
    """Convert markdown paper to formatted Word document."""
    if not HAS_DOCX:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    paper = parse_markdown(md_text)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title_text = paper.title or "论文标题"
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after = Pt(12)
    run = tp.add_run(title_text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    if paper.abstract:
        ap = doc.add_paragraph()
        ap.paragraph_format.space_before = Pt(6)
        ap.paragraph_format.space_after = Pt(6)
        run = ap.add_run("摘要　")
        run.bold = True
        ap.add_run(paper.abstract)

    if paper.keywords:
        kp = doc.add_paragraph()
        kp.paragraph_format.space_after = Pt(12)
        run = kp.add_run("关键词：")
        run.bold = True
        kp.add_run(paper.keywords)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()

    for section_data in paper.sections:
        level = section_data.level
        title = section_data.title

        if level == 1:
            hp = doc.add_heading(title, level=1)
            for run in hp.runs:
                run.font.size = Pt(14)
                run.bold = True
        elif level == 2:
            hp = doc.add_heading(title, level=2)
            for run in hp.runs:
                run.font.size = Pt(12)
                run.bold = True
        elif level == 3:
            hp = doc.add_heading(title, level=3)
            for run in hp.runs:
                run.bold = True
        else:
            hp = doc.add_heading(title, level=4)
            for run in hp.runs:
                run.bold = True

        for item in section_data.content:
            if hasattr(item, "headers"):
                add_threeline_table(doc, {
                    "headers": item.headers,
                    "rows": item.rows,
                    "caption": item.caption
                })
            elif isinstance(item, str):
                render_text_paragraph(doc, item)
            elif isinstance(item, list):
                for sub in item:
                    if isinstance(sub, str):
                        render_text_paragraph(doc, sub)

    if paper.references:
        doc.add_heading("参考文献", level=1)
        for ref in paper.references:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(-0.74)
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(ref)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    doc.save(output_path)
    print(f"[OK] Word document saved: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# LaTeX / PDF Generation — IEEE Transactions Format
# ─────────────────────────────────────────────────────────────────────────────

# Resolved at runtime via get_ieee_template_path()
_ref_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "references")


def get_ieee_template_path():
    """Return the path to the bundled IEEE Transactions template."""
    return os.path.join(_ref_dir, "ieee_transactions_template.tex")


def get_IEEEtran_cls_path():
    """Return the path to IEEEtran.cls."""
    return os.path.join(_ref_dir, "IEEEtran.cls")


# ─── LaTeX helpers ────────────────────────────────────────────────────────────

def escape_latex(text):
    """Escape special LaTeX characters."""
    if not text:
        return ""
    text = str(text)
    for char, replacement in [
        ("&", r"\&"),
        ("%", r"\%"),
        ("$", r"\$"),
        ("#", r"\#"),
        ("_", r"\_"),
        ("{", r"\{"),
        ("}", r"\}"),
        ("~", r"\textasciitilde{}"),
        ("^", r"\textasciicircum{}"),
        ("\\", r"\textbackslash{}"),
    ]:
        text = text.replace(char, replacement)
    return text


def md_to_latex_text(text):
    """Convert markdown inline formatting to LaTeX."""
    # Bold and italic
    text = re.sub(r"\*\*([^*]+)\*\*", r"\\textbf{\1}", text)
    text = re.sub(r"\*([^*]+)\*", r"\\textit{\1}", text)
    # Clean up escaped asterisks that were part of the bold/italic syntax
    text = text.replace(r"\*", "*")
    return text


def md_to_latex_text_preserve_newlines(text):
    """Convert markdown inline + paragraphs to LaTeX, preserving newlines as \\."""
    lines = text.split("\n")
    parts = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            parts.append("")
            continue
        # Escape inline formatting
        segs = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", stripped)
        rendered = []
        for seg in segs:
            if seg.startswith("**") and seg.endswith("**"):
                rendered.append(r"\textbf{" + seg[2:-2] + "}")
            elif seg.startswith("*") and seg.endswith("*"):
                rendered.append(r"\textit{" + seg[1:-1] + "}")
            else:
                rendered.append(escape_latex(seg))
        parts.append(" ".join(rendered))
    return "\n\n".join(parts)


def table_to_latex(table_data):
    """Convert table data to IEEE-style LaTeX three-line table (booktabs)."""
    headers = table_data["headers"]
    rows = table_data["rows"]
    caption = table_data.get("caption", "")

    lines = []
    n_cols = len(headers)
    lines.append(r"\begin{table}[htbp]")
    lines.append(r"  \centering")
    lines.append(r"  \caption{" + escape_latex(caption) + "}" if caption else r"  \caption{}")
    lines.append(r"  \begin{tabular}{" + "c" * n_cols + "}")
    lines.append(r"    \toprule")
    lines.append("    " + " & ".join(escape_latex(h) for h in headers) + r" \\")
    lines.append(r"    \midrule")
    for row in rows:
        cells = [escape_latex(str(c)) for c in row]
        lines.append("    " + " & ".join(cells) + r" \\")
    lines.append(r"    \bottomrule")
    lines.append(r"  \end{tabular}")
    lines.append(r"\end{table}")
    return "\n".join(lines)


def latex_to_text_mode(text):
    """Convert LaTeX source text to plain text (for abstract/keywords fields)."""
    # Remove common LaTeX commands used in abstracts
    text = re.sub(r"\\[a-zA-Z]+\{[^}]*\}", lambda m: m.group(1) if "{" in m.group(0) else "", text)
    text = re.sub(r"\\&", "&", text)
    text = re.sub(r"\\\$", "$", text)
    text = re.sub(r"\\%s?" % "", text)
    return text


# ─── Build IEEE LaTeX document ───────────────────────────────────────────────

def build_ieee_author_block(paper):
    """
    Build IEEE author block with \thanks{} for affiliations.
    Falls back to a simple centered author line if no author info available.
    """
    authors_text = paper.authors.strip() if paper.authors else ""

    if not authors_text or authors_text == "作者":
        # No author info — use a placeholder that won't break compilation
        authors_text = "First Author, Second Author, and Third Author"

    # Split by common separators: "，" (Chinese), ",", " and ", " & "
    import re as _re
    authors = _re.split(r'[，,&]+', authors_text)
    authors = [a.strip() for a in authors if a.strip()]

    if not authors:
        authors = ["Author"]

    lines = []
    lines.append(r"\author{")
    author_entries = []
    for i, author in enumerate(authors):
        author = author.strip()
        if not author:
            continue
        # Wrap each author with a \thanks{} containing a placeholder affiliation
        # \thanks{} content won't break compilation even if left as placeholder
        entry = f"{author}\\thanks{{Author affiliation to be completed.}}"
        author_entries.append(entry)

    lines.append("    " + "\\\\\n    ".join(author_entries))
    lines.append(r"}")
    return "\n".join(lines)


def build_ieee_markboth(paper):
    """Build \markboth{...}{...} for running headers."""
    short_title = paper.title[:60] + "..." if len(paper.title) > 60 else paper.title
    left = "IEEE Transactions on Cybernetics"
    right = f"Author: {short_title}"
    return f"\\markboth{{{left}}}{{{right}}}"


def build_ieee_section_content(section_data, is_first_section=False):
    """
    Convert a Section's content to LaTeX string.
    If is_first_section=True (Introduction), prepend \IEEEPARstart.
    """
    level = section_data.level
    title = section_data.title

    # IEEE section commands: \section, \subsection, \subsubsection
    sec_cmd = {1: "section", 2: "subsection", 3: "subsubsection"}.get(level, "paragraph")
    lines = []
    lines.append(f"\\{{{sec_cmd}}}{{{escape_latex(title)}}}")

    # First paragraph of Introduction gets \IEEEPARstart
    use_parstart = is_first_section

    for item in section_data.content:
        if hasattr(item, "headers"):
            lines.append(table_to_latex({
                "headers": item.headers,
                "rows": item.rows,
                "caption": item.caption
            }))
            lines.append("")
        elif isinstance(item, str):
            para = md_to_latex_text_preserve_newlines(item)
            if para.strip():
                if use_parstart:
                    # Take the first sentence's first letter for \IEEEPARstart
                    first_letter = item.strip()[0] if item.strip() else "T"
                    second_letter = item.strip()[1] if len(item.strip()) > 1 else "he"
                    para = rf"\IEEEPARstart{{{first_letter}}}{{{second_letter}}}{para[len(item.strip()[:2]):].strip()}"
                    use_parstart = False
                lines.append(para)
                lines.append("")
        elif isinstance(item, list):
            for sub in item:
                if isinstance(sub, str):
                    para = md_to_latex_text_preserve_newlines(sub)
                    if para.strip():
                        lines.append(para)
                        lines.append("")
    return "\n".join(lines)


def build_ieee_body(paper):
    """
    Build the full body LaTeX string for IEEE template.
    Splits content into: main sections, appendices (before references).
    """
    body_parts = []
    appendices_parts = []

    main_sections_count = 0

    for section_data in paper.sections:
        if section_data.is_appendix:
            appendices_parts.append(build_ieee_section_content(section_data))
        elif section_data.is_acknowledgment:
            # Acknowledgment is rendered after appendices but before references
            ack_cmd = r"\section*{Acknowledgment}"
            ack_lines = [ack_cmd]
            for item in section_data.content:
                if hasattr(item, "headers"):
                    ack_lines.append(table_to_latex({
                        "headers": item.headers, "rows": item.rows,
                        "caption": item.caption
                    }))
                elif isinstance(item, str):
                    ack_lines.append(md_to_latex_text_preserve_newlines(item))
            appendices_parts.append("\n".join(ack_lines))
        else:
            is_first_main = (main_sections_count == 0)
            body_parts.append(build_ieee_section_content(section_data, is_first_section=is_first_main))
            main_sections_count += 1

    all_body = "\n\n".join(body_parts)

    # Appendices: insert \appendices before them
    if appendices_parts:
        all_body += "\n\n\\appendices\n\n" + "\n\n".join(appendices_parts)

    return all_body


def build_ieee_bibliography(paper):
    """Build IEEE-style thebibliography section."""
    if not paper.references:
        return ""

    lines = []
    lines.append(r"\IEEEtriggeratref{1}")
    lines.append(r"\IEEEtriggercmd{\textbf{}}")
    lines.append(r"\begin{thebibliography}{99}")

    for ref in paper.references:
        # Convert markdown-style [1] reference to \bibitem
        ref_text = escape_latex(ref)
        # Try to extract the citation key
        key_match = re.match(r'\[(\d+)\]\s*(.+)', ref_text)
        if key_match:
            bib_num = key_match.group(1)
            bib_text = key_match.group(2).strip()
        else:
            bib_text = ref_text
            bib_num = "?"
        idx = paper.references.index(ref) + 1
        lines.append("\\bibitem{[ref" + str(idx) + "]}")
        lines.append("  " + bib_text)
    lines.append(r"\end{thebibliography}")
    return "\n".join(lines)


def paper_to_ieee_latex(paper, template_path=None):
    """
    Convert a parsed Paper to IEEE Transactions LaTeX document string.
    Uses ieee_transactions_template.tex as the base.
    """
    if template_path is None:
        template_path = get_ieee_template_path()

    with open(template_path, "r", encoding="utf-8") as f:
        template = f.read()

    # ── Title & Author block ─────────────────────────────────────────────────
    title_escaped = escape_latex(paper.title) if paper.title else "Paper Title"
    author_block = build_ieee_author_block(paper)
    markboth = build_ieee_markboth(paper)

    # ── Abstract ─────────────────────────────────────────────────────────────
    abstract_text = paper.abstract.strip() if paper.abstract else "Abstract to be completed."

    # ── Keywords ─────────────────────────────────────────────────────────────
    keywords_text = paper.keywords.strip() if paper.keywords else "keyword1, keyword2, keyword3"

    # ── Body ─────────────────────────────────────────────────────────────────
    body_latex = build_ieee_body(paper)

    # ── Bibliography ───────────────────────────────────────────────────────────
    bibliography_latex = build_ieee_bibliography(paper)

    # ── Substitute placeholders ──────────────────────────────────────────────
    replacements = [
        ("__TITLE__", title_escaped),
        ("__AUTHORS__", author_block + "\n" + markboth),
        ("__ABSTRACT__", escape_latex(abstract_text)),
        ("__KEYWORDS__", escape_latex(keywords_text)),
        ("__BODY__", body_latex),
        ("__BIBLIOGRAPHY__", bibliography_latex),
    ]

    result = template
    for placeholder, value in replacements:
        result = result.replace(placeholder, value)

    return result


def compile_latex_to_pdf(tex_path, output_dir, cls_dir=None):
    """
    Compile LaTeX to PDF using xelatex.
    Copies IEEEtran.cls/bst to the tex directory if provided.
    """
    basename = os.path.splitext(os.path.basename(tex_path))[0]
    pdf_path = os.path.join(output_dir, basename + ".pdf")
    tex_dir = os.path.dirname(os.path.abspath(tex_path))

    # Copy IEEEtran files alongside the tex file if they exist in references
    ref_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "references")
    for fname in ["IEEEtran.cls", "IEEEtran.bst"]:
        src = os.path.join(ref_dir, fname)
        if os.path.exists(src):
            shutil.copy(src, tex_dir)

    # Try xelatex first (better for Unicode/Chinese), fall back to pdflatex
    engines = ["xelatex", "pdflatex"]
    compiled = False

    for engine in engines:
        # Clean up old auxiliary files
        for ext in [".aux", ".log", ".out", ".bbl", ".blg", ".toc", ".fls"]:
            aux_file = os.path.join(output_dir, basename + ext)
            if os.path.exists(aux_file):
                os.remove(aux_file)

        for pass_num in range(2):
            cmd = [
                engine,
                "-interaction=nonstopmode",
                "-halt-on-error",
                f"-output-directory={output_dir}",
                tex_path
            ]
            result = subprocess.run(
                cmd, capture_output=True, text=True, timeout=120
            )
            if result.returncode != 0:
                errors = [
                    l for l in result.stderr.split("\n")
                    if l.strip() and ("Error" in l.upper() or "!" in l or "Fatal" in l)
                ]
                if errors:
                    last_error = "\n".join(errors[:5])
                    if engine == engines[-1]:  # last engine, last pass
                        print(f"[ERROR] {engine} compilation failed:\n{last_error}")
                break
        else:
            # All passes succeeded (break not triggered)
            if os.path.exists(pdf_path):
                compiled = True
                print(f"[OK] PDF saved: {pdf_path}")
                break

    if not compiled and not os.path.exists(pdf_path):
        # Try one more pass with biber for bibliography
        for engine in ["xelatex", "pdflatex"]:
            for pass_num in range(3):
                cmd = [
                    engine,
                    "-interaction=nonstopmode",
                    "-halt-on-error",
                    f"-output-directory={output_dir}",
                    tex_path
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
                if result.returncode == 0:
                    break
            if os.path.exists(pdf_path):
                compiled = True
                break

    if not os.path.exists(pdf_path):
        raise RuntimeError(
            f"PDF compilation failed. Check LaTeX errors above.\n"
            f"Make sure 'xelatex' or 'pdflatex' is installed "
            "(TeX Live: sudo apt install texlive-xetex)"
        )

    return pdf_path


def markdown_to_pdf(md_text, output_path, template_tex=None):
    """
    Convert markdown paper to IEEE Transactions PDF via LaTeX.
    If template_tex is not provided, uses the bundled IEEE template.
    """
    paper = parse_markdown(md_text)

    if template_tex and os.path.exists(template_tex):
        template_path = template_tex
        print(f"[INFO] Using custom LaTeX template: {template_tex}")
    else:
        template_path = get_ieee_template_path()
        print(f"[INFO] Using IEEE Transactions template: {template_path}")

    tex_content = paper_to_ieee_latex(paper, template_path=template_path)

    tex_dir = os.path.dirname(os.path.abspath(output_path)) or "."
    tex_basename = os.path.splitext(os.path.basename(output_path))[0]
    tex_path = os.path.join(tex_dir, tex_basename + ".tex")
    with open(tex_path, "w", encoding="utf-8") as f:
        f.write(tex_content)
    print(f"[INFO] LaTeX file written: {tex_path}")

    # Copy IEEEtran files to the output directory for compilation
    ref_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "references")
    for fname in ["IEEEtran.cls", "IEEEtran.bst"]:
        src = os.path.join(ref_dir, fname)
        if os.path.exists(src):
            shutil.copy(src, tex_dir)

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_tex = os.path.join(tmp_dir, tex_basename + ".tex")
        shutil.copy(tex_path, tmp_tex)
        # Also copy IEEEtran files to tmp dir
        for fname in ["IEEEtran.cls", "IEEEtran.bst"]:
            src = os.path.join(ref_dir, fname)
            if os.path.exists(src):
                shutil.copy(src, tmp_dir)
        pdf_path = compile_latex_to_pdf(tmp_tex, tmp_dir)
        final_pdf = os.path.join(tex_dir, tex_basename + ".pdf")
        shutil.copy(pdf_path, final_pdf)

    print(f"[OK] PDF saved: {final_pdf}")
    return final_pdf


# ─────────────────────────────────────────────────────────────────────────────
# Main entry point
# ─────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown paper to Word (.docx) or IEEE PDF"
    )
    parser.add_argument("--input", required=True, help="Input markdown file path")
    parser.add_argument(
        "--output_format", required=True, choices=["docx", "pdf"],
        help="Output format: docx or pdf"
    )
    parser.add_argument("--output", required=True, help="Output file path (with extension)")
    parser.add_argument("--reference", help="Path to reference .docx template (for Word output)")
    parser.add_argument(
        "--template",
        help="Path to custom LaTeX template .tex (for PDF output). "
             "If omitted, uses the bundled IEEE Transactions template."
    )

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"[ERROR] Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    with open(args.input, "r", encoding="utf-8") as f:
        md_text = f.read()

    if args.output_format == "docx":
        markdown_to_docx(md_text, args.output, ref_template=args.reference)
    elif args.output_format == "pdf":
        markdown_to_pdf(md_text, args.output, template_tex=args.template)
    else:
        print(f"[ERROR] Unknown format: {args.output_format}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
