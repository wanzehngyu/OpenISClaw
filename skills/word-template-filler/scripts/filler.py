#!/usr/bin/env python3
"""
Word Template Filler: Fill a Word template with markdown paper content.

Usage:
    python3 filler.py --template TEMPLATE.docx --paper PAPER.md --output OUTPUT.docx [--pdf]
"""

import argparse
import os
import re
import subprocess
import sys
import shutil
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)

from md_parser import parse_markdown_paper, Paper


# ─────────────────────────────────────────────────────────────────────────────
# Table creation (three-line academic style)
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set borders for a table cell. Each arg is dict(val, sz, color) or None for nil."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, spec in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        border = OxmlElement(f"w:{edge}")
        if spec:
            border.set(qn("w:val"), spec.get("val", "single"))
            border.set(qn("w:sz"), str(spec.get("sz", 4)))
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), spec.get("color", "000000"))
        else:
            border.set(qn("w:val"), "nil")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_three_line_table(doc, headers, rows, caption=""):
    """Add a three-line academic table to the document."""
    n_cols = len(headers)
    table = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)
        set_cell_border(cell,
            top={"val": "double", "sz": 6, "color": "000000"},
            bottom={"val": "single", "sz": 4, "color": "000000"},
            left={"val": "nil"},
            right={"val": "nil"},
        )

    # Data rows
    for r_idx, row in enumerate(rows):
        is_last = (r_idx == len(rows) - 1)
        row_el = table.rows[r_idx + 1]
        for j, cell_text in enumerate(row):
            cell = row_el.cells[j]
            cell.text = str(cell_text) if cell_text else ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(10)
            left_spec = {"val": "single", "sz": 4, "color": "000000"} if j == 0 else {"val": "nil"}
            right_spec = {"val": "single", "sz": 4, "color": "000000"} if j == n_cols - 1 else {"val": "nil"}
            bottom_spec = {"val": "double", "sz": 6, "color": "000000"} if is_last else {"val": "single", "sz": 4, "color": "000000"}
            set_cell_border(cell,
                top={"val": "nil"},
                bottom=bottom_spec,
                left=left_spec,
                right=right_spec,
            )

    # Caption
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after = Pt(6)
        run = cp.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)

    return table


# ─────────────────────────────────────────────────────────────────────────────
# Core placeholder replacement
# ─────────────────────────────────────────────────────────────────────────────

def find_placeholder_paragraphs(doc):
    """Find all paragraphs containing template placeholders."""
    results = []
    for para in doc.paragraphs:
        for run in para.runs:
            # Match {{NAME}} pattern (double braces)
            found = re.findall(r"\{\{[^}]+\}\}", run.text)
            for ph in found:
                results.append((para, ph, run))
    return results


def replace_placeholder_text(para, placeholder, new_text, preserve_format=True):
    """Replace a {{PLACEHOLDER}} in a paragraph with new text."""
    target_run = None
    for run in para.runs:
        if placeholder in run.text:
            target_run = run
            break

    if target_run is None:
        return False

    fmt = {
        "bold": target_run.bold,
        "italic": target_run.italic,
        "underline": target_run.underline,
        "font_name": target_run.font.name,
        "font_size": target_run.font.size,
    }

    target_run.text = target_run.text.replace(placeholder, new_text, 1)

    if preserve_format:
        if fmt["bold"] is not None:
            target_run.bold = fmt["bold"]
        if fmt["italic"] is not None:
            target_run.italic = fmt["italic"]
        if fmt["underline"]:
            target_run.underline = fmt["underline"]
        if fmt["font_name"]:
            target_run.font.name = fmt["font_name"]
        if fmt["font_size"]:
            target_run.font.size = fmt["font_size"]

    return True


def replace_placeholder_with_table(doc, para, placeholder, table_data):
    """Replace a table placeholder with an actual three-line table."""
    # Save paragraph format info
    para_fmt = para.paragraph_format

    # Insert table after this paragraph
    tbl = add_three_line_table(doc, table_data.headers, table_data.rows, table_data.caption)

    # Remove the placeholder text from paragraph
    for run in para.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, "", 1)
            if not run.text.strip():
                # Mark run for removal
                pass
            break


def replace_placeholder_with_refs(doc, para, placeholder, ref_list):
    """Replace a ref placeholder with hanging-indent references."""
    if not ref_list:
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, "", 1)
                break
        return

    # Replace in first paragraph
    first_ref = ref_list[0]
    for run in para.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, first_ref, 1)
            # Apply formatting
            run.font.size = Pt(12)
            # Apply hanging indent
            para.paragraph_format.first_line_indent = Cm(-0.74)
            para.paragraph_format.left_indent = Cm(0.74)
            para.paragraph_format.space_after = Pt(3)
            break

    # Add remaining refs as new paragraphs
    last_para = para
    for ref in ref_list[1:]:
        new_para = doc.add_paragraph()
        new_para.paragraph_format.first_line_indent = Cm(-0.74)
        new_para.paragraph_format.left_indent = Cm(0.74)
        new_para.paragraph_format.space_after = Pt(3)
        new_run = new_para.add_run(ref)
        new_run.font.size = Pt(12)
        last_para = new_para


def fill_template(template_path: str, paper: Paper, output_path: str):
    """Fill a Word template with paper content."""
    doc = Document(template_path)

    # Build flat list of all content blocks (paragraphs and tables) in order
    all_content_blocks: list = []
    for section in paper.sections:
        for item in section.content:
            all_content_blocks.append(item)

    # Count tables for TABLE:N lookup
    paper_tables = [block.table for block in all_content_blocks if block.table is not None]
    paper_paras = [block for block in all_content_blocks if block.text is not None]

    table_block_idx = 0
    para_block_idx = 0

    # Process each placeholder
    placeholder_paras = find_placeholder_paragraphs(doc)
    print(f"[INFO] Found {len(placeholder_paras)} placeholders")

    for para, placeholder, run in placeholder_paras:
        ph = placeholder.strip()
        ph_upper = ph.upper()
        replaced = False

        # TITLE
        if ph_upper == "{{TITLE}}":
            replaced = replace_placeholder_text(para, placeholder, paper.title)

        # ABSTRACT
        elif ph_upper == "{{ABSTRACT}}":
            replaced = replace_placeholder_text(para, placeholder, paper.abstract)

        # KEYWORDS
        elif ph_upper == "{{KEYWORDS}}":
            replaced = replace_placeholder_text(para, placeholder, paper.keywords)

        # SECTION:N → fill section heading
        elif "{{SECTION:" in ph:
            m = re.search(r"\{\{SECTION:(\d+)\}\}", ph, re.IGNORECASE)
            if m:
                idx = int(m.group(1)) - 1
                if idx < len(paper.sections):
                    sec = paper.sections[idx]
                    heading_text = f"{sec.number}  {sec.title}"
                    replaced = replace_placeholder_text(para, placeholder, heading_text)
                    # Force bold for section headings
                    for r in para.runs:
                        if placeholder in r.text:
                            r.text = r.text.replace(placeholder, heading_text, 1)
                            r.bold = True
                            replaced = True
                            break

        # TABLE:N → insert three-line table
        elif "{{TABLE:" in ph:
            m = re.search(r"\{\{TABLE:(\d+)\}\}", ph, re.IGNORECASE)
            if m:
                idx = int(m.group(1)) - 1
                if idx < len(paper_tables):
                    replace_placeholder_with_table(doc, para, placeholder, paper_tables[idx])
                    replaced = True
                    table_block_idx += 1

        # CONTENT:N → fill paragraph text
        elif "{{CONTENT:" in ph:
            m = re.search(r"\{\{CONTENT:(\d+)\}\}", ph, re.IGNORECASE)
            if m:
                idx = int(m.group(1)) - 1
                if idx < len(paper_paras):
                    replaced = replace_placeholder_text(para, placeholder, paper_paras[idx].text)

        # REF → references
        elif ph_upper in ("{{REF}}", "{{参考文献}}"):
            replace_placeholder_with_refs(doc, para, placeholder, paper.references)
            replaced = True

        else:
            print(f"[WARN] Unknown placeholder: {ph}")

    # Save output
    output_dir = os.path.dirname(os.path.abspath(output_path)) or "."
    temp_doc = os.path.join(output_dir, os.path.basename(output_path))
    if not temp_doc.lower().endswith(".docx"):
        temp_doc += ".docx"
    doc.save(temp_doc)
    print(f"[OK] Document saved: {temp_doc}")
    return temp_doc


def convert_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Convert Word document to PDF using LibreOffice."""
    lo_paths = [
        "/opt/homebrew/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "soffice",
    ]
    lo_cmd = next((p for p in lo_paths if os.path.exists(p) or p == "soffice"), None)

    if lo_cmd is None:
        print("[WARN] LibreOffice not found. Install: brew install libreoffice")
        return False

    output_dir = os.path.dirname(os.path.abspath(pdf_path)) or "."
    base = os.path.splitext(os.path.basename(docx_path))[0]
    cmd = [lo_cmd, "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            print(f"[WARN] LibreOffice error: {result.stderr[:200]}")
            return False
        generated_pdf = os.path.join(output_dir, base + ".pdf")
        if os.path.exists(generated_pdf) and generated_pdf != pdf_path:
            shutil.copy(generated_pdf, pdf_path)
        print(f"[OK] PDF saved: {pdf_path}")
        return True
    except subprocess.TimeoutExpired:
        print("[WARN] LibreOffice timed out")
        return False
    except Exception as e:
        print(f"[WARN] PDF conversion error: {e}")
        return False


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Fill Word template with markdown paper content")
    parser.add_argument("--template", required=True)
    parser.add_argument("--paper", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--pdf", action="store_true")

    args = parser.parse_args()

    if not os.path.exists(args.template):
        print(f"[ERROR] Template not found: {args.template}")
        sys.exit(1)
    if not os.path.exists(args.paper):
        print(f"[ERROR] Paper not found: {args.paper}")
        sys.exit(1)

    print(f"[INFO] Parsing paper: {args.paper}")
    with open(args.paper, "r", encoding="utf-8") as f:
        paper = parse_markdown_paper(f.read())
    print(f"  Title: {paper.title}")
    print(f"  Sections: {len(paper.sections)}")
    tables = sum(1 for s in paper.sections for c in s.content if c.table)
    print(f"  Tables: {tables}")
    print(f"  References: {len(paper.references)}")

    output_docx = args.output
    if not output_docx.lower().endswith(".docx"):
        output_docx += ".docx"

    print(f"[INFO] Filling template...")
    filled = fill_template(args.template, paper, output_docx)

    if args.pdf:
        pdf_path = os.path.splitext(output_docx)[0] + ".pdf"
        convert_to_pdf(filled, pdf_path)

    print("[OK] Done!")


if __name__ == "__main__":
    main()
