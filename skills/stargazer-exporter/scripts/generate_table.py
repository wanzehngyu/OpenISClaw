#!/usr/bin/env python3
"""
Stargazer Exporter: Academic Table Formatter for Econometric Results

Usage:
    python generate_table.py --pickles <path1> [path2...]]
                              [--models <name1> <name2>...]
                              [--rename <old1:new1,old2:new2>]
                              [--title <table_title>]
                              --output_dir <path>
                              [--formats latex,html,docx]
"""

import sys
import os
import pickle
import argparse
import warnings
from pathlib import Path

import pandas as pd
import numpy as np


# Default rename map for common IS econometric variables
DEFAULT_RENAME_MAP = {
    "it_investment_g": "IT Investment (%)",
    "co_size_ln": "Firm Size (log)",
    "lev": "Leverage Ratio",
    "roa": "Return on Assets",
    "roe": "Return on Equity",
    "predicted_it": "IT Investment (IV)",
    "ln_gov_proc": "Government Procurement (log)",
    "digital_infrastructure": "Digital Infrastructure",
    "age": "Firm Age",
    "tfp_lp": "TFP (LP)",
    "labour_productivity": "Labour Productivity",
    "export_intensity": "Export Intensity",
    "rd_intensity": "R&D Intensity",
    "hhi": "Industry Concentration (HHI)",
    "ceo_tenure": "CEO Tenure",
    "board_size": "Board Size",
}


def parse_rename_map(rename_str):
    """Parse rename mapping from string format 'old1:new1,old2:new2'."""
    if not rename_str:
        return {}
    rename_map = {}
    for pair in rename_str.split(","):
        if ":" in pair:
            old, new = pair.split(":", 1)
            rename_map[old.strip()] = new.strip()
    return rename_map


def load_pickle_results(pickle_paths):
    """Load and validate pickle files."""
    results = []
    for path in pickle_paths:
        if not os.path.exists(path):
            print(f"⚠️ [警告] 文件不存在: {path}，跳过。")
            continue
        try:
            with open(path, "rb") as f:
                data = pickle.load(f)
                # Handle dict wrapper (from iv-estimator)
                if isinstance(data, dict) and "results" in data:
                    results.append(data["results"])
                # Handle raw results
                else:
                    results.append(data)
        except Exception as e:
            print(f"⚠️ [警告] 无法读取 {path}: {str(e)}，跳过。")
            continue
    return results


def get_model_summary(result):
    """Extract summary info from a regression result."""
    summary = {
        "coefficients": {},
        "std_errors": {},
        "pvalues": {},
        "n_obs": None,
        "r_squared": None,
        "f_statistic": None,
    }

    try:
        # linearmodels results
        if hasattr(result, "summary"):
            # PanelOLS / IV2SLS results
            params = result.summary.params
            se = result.summary.std_errors
            pvalues = result.summary.pvalues

            for i, (var, coef) in enumerate(params.items()):
                summary["coefficients"][var] = coef
                summary["std_errors"][var] = se[i] if i < len(se) else None
                summary["pvalues"][var] = pvalues[i] if i < len(pvalues) else None

            summary["n_obs"] = getattr(result, "nobs", None)
            summary["r_squared"] = getattr(result, "rsquared_within", None) or getattr(result, "rsquared", None)
            summary["f_statistic"] = getattr(result, "f_statistic", None)
            if summary["f_statistic"] is not None and hasattr(summary["f_statistic"], "value"):
                summary["f_statistic"] = summary["f_statistic"].value

        elif hasattr(result, "params"):
            # Standard statsmodels-style results
            params = result.params
            for var, coef in params.items():
                summary["coefficients"][var] = coef
                summary["std_errors"][var] = result.std_errors.get(var, None) if hasattr(result, "std_errors") else None
                summary["pvalues"][var] = result.pvalues.get(var, None) if hasattr(result, "pvalues") else None

            summary["n_obs"] = getattr(result, "nobs", None)
            summary["r_squared"] = getattr(result, "rsquared", None)
            summary["f_statistic"] = getattr(result, "fvalue", None)

    except Exception as e:
        print(f"⚠️ [警告] 结果解析失败: {str(e)}")

    return summary


def format_coef(coef, se, pvalue):
    """Format coefficient with significance stars."""
    if coef is None or se is None:
        return "N/A", ""

    coef_val = float(coef)
    se_val = float(se)

    # Determine significance
    if pvalue is not None:
        pval = float(pvalue)
        if pval < 0.01:
            stars = "***"
        elif pval < 0.05:
            stars = "**"
        elif pval < 0.1:
            stars = "*"
        else:
            stars = ""
    else:
        stars = ""

    # Format with standard errors in parentheses below
    coef_str = f"{coef_val:.4f}"
    se_str = f"({se_val:.4f})"

    return coef_str + stars, se_str


def generate_latex_table(results, model_names, rename_map, title):
    """Generate LaTeX table code."""
    if not results:
        return None

    # Get all variable names from first model
    first_vars = list(results[0].get("coefficients", {}).keys())
    n_models = len(results)

    # Build header
    header = "\\hline\n"
    header += " & ".join(["", *[f"({i+1}) {model_names[i]}" for i in range(n_models)]]) + " \\\\\n"
    header += "\\hline\n"

    # Build rows
    rows = []
    for var in first_vars:
        rename = rename_map.get(var, var)
        row = [rename]

        for res in results:
            coef_dict = res.get("coefficients", {})
            se_dict = res.get("std_errors", {})
            pval_dict = res.get("pvalues", {})

            coef_str, se_str = format_coef(
                coef_dict.get(var),
                se_dict.get(var),
                pval_dict.get(var)
            )
            row.append(coef_str + " \\\\ \n                " + se_str)

        rows.append(" & ".join(row) + " \\\\")
        rows.append("                ")

    # Footer
    footer = "\\hline\n"
    footer += "样本量             & " + " & ".join([str(r.get("n_obs", "N/A")) for r in results]) + " \\\\\n"
    r2_vals = []
    for r in results:
        r2 = r.get("r_squared")
        r2_vals.append(f"{r2:.4f}" if r2 is not None else "N/A")
    footer += "R²                 & " + " & ".join(r2_vals) + " \\\\\n"
    footer += "\\hline\n"

    latex = f"""\\begin{{table}}[htbp]
  \\centering
  \\caption{{{title}}}
  \\begin{{tabular}}{{l{"c" * n_models}}}
    \\hline
    & \\multicolumn{{{n_models}}}{{c}}{{模型}} \\\\
    \\cline{{2-{n_models + 1}}}
    \\hline
    {header}
    {chr(10).join(rows)}
    {footer}
    \\multicolumn{{{n_models + 1}}}{{l}}{{\\footnotesize 括号内为聚类稳健标准误；*** p<0.01, ** p<0.05, * p<0.1}} \\\\
  \\end{{tabular}}
\\end{{table}}"""

    return latex


def generate_html_table(results, model_names, rename_map, title):
    """Generate HTML table."""
    if not results:
        return None

    first_vars = list(results[0].get("coefficients", {}).keys())
    n_models = len(results)

    html = f"""<table border="1" cellpadding="4" cellspacing="0">
  <caption>{title}</caption>
  <thead>
    <tr>
      <th>变量</th>
      {"".join([f'<th colspan="2">({i+1}) {model_names[i]}</th>' for i in range(n_models)])}
    </tr>
  </thead>
  <tbody>
"""

    for var in first_vars:
        rename = rename_map.get(var, var)
        html += f"    <tr><td rowspan=\"2\">{rename}</td>"
        for res in results:
            coef_dict = res.get("coefficients", {})
            se_dict = res.get("std_errors", {})
            pval_dict = res.get("pvalues", {})
            coef_str, se_str = format_coef(
                coef_dict.get(var),
                se_dict.get(var),
                pval_dict.get(var)
            )
            html += f"<td>{coef_str}</td><td>{se_str}</td>"
        html += "</tr>\n"

    # Footer
    html += "    <tr><td colspan=\"" + str(n_models * 2 + 1) + "\">"
    html += f"<small>样本量: {results[0].get('n_obs', 'N/A')} | R²: {results[0].get('r_squared', 'N/A'):.4f if results[0].get('r_squared') else 'N/A'}</small>"
    html += "</td></tr>\n"

    html += "  </tbody>\n</table>"
    return html


def generate_ascii_table(results, model_names, rename_map, title):
    """Generate ASCII table."""
    if not results:
        return None

    n_models = len(results)
    col_width = 20

    # Header
    header = f"{'变量':<20}" + "".join([f"{(i+1) + model_names[i]:^40}" for i in range(n_models)])
    sep = "-" * (20 + 40 * n_models)

    rows = []
    first_vars = list(results[0].get("coefficients", {}).keys())
    for var in first_vars:
        rename = rename_map.get(var, var)
        row = f"{rename:<20}"
        for res in results:
            coef_dict = res.get("coefficients", {})
            se_dict = res.get("std_errors", {})
            pval_dict = res.get("pvalues", {})
            coef_str, se_str = format_coef(
                coef_dict.get(var),
                se_dict.get(var),
                pval_dict.get(var)
            )
            row += f"{coef_str + ' / ' + se_str:^40}"
        rows.append(row)

    # Footer
    footer_row = f"{'样本量':<20}"
    footer_row += "".join([f"{results[i].get('n_obs', 'N/A'):^40}" for i in range(n_models)])

    ascii_table = f"""
{title}
{'=' * (20 + 40 * n_models)}
{header}
{sep}
{chr(10).join(rows)}
{sep}
{footer_row}
{'=' * (20 + 40 * n_models)}
* Standard errors in parentheses
*** p<0.01, ** p<0.05, * p<0.1
"""
    return ascii_table


def main():
    parser = argparse.ArgumentParser(description="Stargazer: Academic Table Generator")
    parser.add_argument("--pickles", required=True, nargs="+", help="Pickle result file paths")
    parser.add_argument("--models", nargs="*", default=[], help="Model names for column headers")
    parser.add_argument("--rename", default="", help="Rename mapping: var1:label1,var2:label2,...")
    parser.add_argument("--title", default="实证计量回归分析结果", help="Table title")
    parser.add_argument("--output_dir", required=True, help="Output directory")
    parser.add_argument("--formats", default="latex", help="Output formats: latex,html,ascii,docx (comma-separated)")

    args = parser.parse_args()

    # Parse formats
    formats = [f.strip() for f in args.formats.split(",")]
    rename_map = {**DEFAULT_RENAME_MAP, **parse_rename_map(args.rename)}

    # Load results
    results = load_pickle_results(args.pickles)
    if not results:
        print("❌ [错误] 未找到有效的回归结果文件。")
        sys.exit(1)

    # Generate model names if not provided
    if not args.models:
        model_names = [f"Model {i+1}" for i in range(len(results))]
    else:
        model_names = args.models[:len(results)]

    # Create output directory
    os.makedirs(args.output_dir, exist_ok=True)

    # Generate outputs
    outputs = {}
    if "latex" in formats:
        latex = generate_latex_table(results, model_names, rename_map, args.title)
        if latex:
            path = os.path.join(args.output_dir, "regression_table.tex")
            with open(path, "w", encoding="utf-8") as f:
                f.write(latex)
            outputs["latex"] = path
            print(f"✅ [LaTeX 表格已保存] {path}")

    if "html" in formats:
        html = generate_html_table(results, model_names, rename_map, args.title)
        if html:
            path = os.path.join(args.output_dir, "regression_table.html")
            with open(path, "w", encoding="utf-8") as f:
                f.write(html)
            outputs["html"] = path
            print(f"✅ [HTML 表格已保存] {path}")

    if "ascii" in formats:
        ascii_table = generate_ascii_table(results, model_names, rename_map, args.title)
        if ascii_table:
            path = os.path.join(args.output_dir, "regression_table.txt")
            with open(path, "w", encoding="utf-8") as f:
                f.write(ascii_table)
            outputs["ascii"] = path
            print(f"✅ [ASCII 表格已保存] {path}")

    if "docx" in formats:
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            title_para = doc.add_heading(args.title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            n_rows = len(results[0].get("coefficients", {})) + 4  # vars + 3 meta rows
            n_cols = 2 * len(results) + 1  # var col + 2 per model (coef, se)

            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.style = "Table Grid"

            # Header row
            table.cell(0, 0).text = "变量"
            for m_idx, model_name in enumerate(model_names):
                coef_cell = table.cell(0, m_idx * 2 + 1)
                se_cell = table.cell(0, m_idx * 2 + 2)
                coef_cell.text = f"({m_idx + 1}) {model_name}"
                se_cell.text = ""
                # Merge header cells
                coef_cell.merge(se_cell)

            # Data rows
            first_vars = list(results[0].get("coefficients", {}).keys())
            for v_idx, var in enumerate(first_vars):
                rename = rename_map.get(var, var)
                row = table.rows[v_idx + 1]
                row.cells[0].text = rename

                for m_idx, res in enumerate(results):
                    coef_dict = res.get("coefficients", {})
                    se_dict = res.get("std_errors", {})
                    pval_dict = res.get("pvalues", {})
                    coef_str, se_str = format_coef(
                        coef_dict.get(var),
                        se_dict.get(var),
                        pval_dict.get(var)
                    )
                    row.cells[m_idx * 2 + 1].text = coef_str
                    row.cells[m_idx * 2 + 2].text = se_str

            # Meta rows
            meta_row_start = len(first_vars) + 1

            # N row
            n_row = table.rows[meta_row_start]
            n_row.cells[0].text = "样本量"
            for m_idx, res in enumerate(results):
                n_row.cells[m_idx * 2 + 1].text = str(res.get("n_obs", "N/A"))
                n_row.cells[m_idx * 2 + 2].text = ""

            # R² row
            r2_row = table.rows[meta_row_start + 1]
            r2_row.cells[0].text = "R²"
            for m_idx, res in enumerate(results):
                r2_val = res.get("r_squared")
                r2_row.cells[m_idx * 2 + 1].text = f"{r2_val:.4f}" if r2_val is not None else "N/A"
                r2_row.cells[m_idx * 2 + 2].text = ""

            # Note row
            note_row = table.rows[meta_row_start + 2]
            note_cell = note_row.cells[0]
            for c_idx in range(1, n_cols):
                note_cell = note_cell.merge(note_row.cells[c_idx])
            note_cell.text = "括号内为聚类稳健标准误；*** p<0.01, ** p<0.05, * p<0.1"

            path = os.path.join(args.output_dir, "regression_table.docx")
            doc.save(path)
            outputs["docx"] = path
            print(f"✅ [Word 表格已保存] {path}")

        except ImportError:
            print("⚠️ [警告] python-docx 未安装，Word 输出跳过。运行: pip install python-docx")

    # Summary
    print(f"\n📋 [生成完成]")
    print(f"   - 模型数量: {len(results)}")
    print(f"   - 输出文件: {len(outputs)}")
    for fmt, path in outputs.items():
        print(f"   - {fmt}: {path}")


if __name__ == "__main__":
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        main()